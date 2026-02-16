"""
extractors.py — File extraction pipeline for Cane.

Supported formats:
  - PDF  (via PyMuPDF/fitz)
  - DOCX (via python-docx)
  - XLSX / CSV (via openpyxl / csv)
  - Images (OCR placeholder — stores as image for CLIP)

Each extractor returns an ExtractionResult with chunks and/or images.
"""
import csv
import re
from pathlib import Path
from typing import Optional

from models import ExtractionResult, Chunk, Image, make_slug
from smart_chunker import smart_chunk_text
from config import CHUNK_SIZE, CHUNK_OVERLAP, EXT_MAP


def extract(filepath: str, extracted_dir: str) -> ExtractionResult:
    """
    Main entry point. Detect file type and delegate to the right extractor.

    Args:
        filepath: Path to the source file
        extracted_dir: Directory for extracted images/frames

    Returns:
        ExtractionResult with chunks and images
    """
    path = Path(filepath)
    ext = path.suffix.lower()
    file_type = EXT_MAP.get(ext, "")

    if not path.exists():
        return ExtractionResult(
            source_file=path.name, source_type=file_type,
            error=f"File not found: {filepath}"
        )

    try:
        if file_type == "pdf":
            return _extract_pdf(path, extracted_dir)
        elif file_type == "docx":
            return _extract_docx(path, extracted_dir)
        elif file_type in ("xlsx", "csv"):
            return _extract_tabular(path, file_type, extracted_dir)
        elif file_type == "image":
            return _extract_image(path, extracted_dir)
        elif file_type in ("audio", "video"):
            return ExtractionResult(
                source_file=path.name, source_type=file_type,
                error="Audio/video extraction disabled for MVP. Install whisper to enable."
            )
        else:
            return ExtractionResult(
                source_file=path.name, source_type=ext,
                error=f"Unsupported file type: {ext}"
            )
    except Exception as e:
        return ExtractionResult(
            source_file=path.name, source_type=file_type,
            error=f"Extraction failed: {str(e)}"
        )


# ═══════════════════════════════════════════════════════════
#  PDF Extractor (PyMuPDF)
# ═══════════════════════════════════════════════════════════

def _extract_pdf(path: Path, extracted_dir: str) -> ExtractionResult:
    import fitz  # PyMuPDF

    doc = fitz.open(str(path))
    slug = make_slug(path.name)
    img_dir = Path(extracted_dir) / slug / "images"
    img_dir.mkdir(parents=True, exist_ok=True)

    all_text = []
    images = []
    page_texts = {}

    for page_num, page in enumerate(doc, start=1):
        # Extract text
        text = page.get_text("text").strip()
        if text:
            page_texts[page_num] = text
            all_text.append(text)

        # Extract images
        for img_idx, img_info in enumerate(page.get_images(full=True)):
            try:
                xref = img_info[0]
                base_image = doc.extract_image(xref)
                if not base_image:
                    continue

                img_bytes = base_image["image"]
                img_ext = base_image.get("ext", "png")
                width = base_image.get("width", 0)
                height = base_image.get("height", 0)

                # Skip tiny images (icons, bullets, etc.)
                if width < 100 or height < 100:
                    continue

                img_filename = f"p{page_num}_{img_idx}.{img_ext}"
                img_path = img_dir / img_filename
                with open(img_path, "wb") as f:
                    f.write(img_bytes)

                images.append(Image(
                    path=str(img_path),
                    source_file=path.name,
                    source_type="pdf",
                    page=page_num,
                    width=width,
                    height=height,
                ))
            except Exception:
                continue

    # OCR fallback: if no text was extracted but PDF has pages, it's likely scanned
    if not page_texts and doc.page_count > 0:
        print(f"  [OCR] No text found in {path.name}, running OCR on {doc.page_count} pages...")
        page_texts = _ocr_pdf_pages(doc, path.name)
        if page_texts:
            print(f"  [OCR] Extracted text from {len(page_texts)} pages")
        else:
            print(f"  [OCR] No text recovered from OCR")

    doc.close()

    # Chunk the text with page awareness
    chunks = _chunk_pages(page_texts, path.name, "pdf")

    print(f"  [PDF] {path.name}: {len(chunks)} chunks, {len(images)} images from {len(page_texts)} pages")

    return ExtractionResult(
        source_file=path.name,
        source_type="pdf",
        chunks=chunks,
        images=images,
    )


# ═══════════════════════════════════════════════════════════
#  DOCX Extractor
# ═══════════════════════════════════════════════════════════

def _extract_docx(path: Path, extracted_dir: str) -> ExtractionResult:
    from docx import Document

    doc = Document(str(path))
    full_text = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            full_text.append(text)

    # Also get tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                full_text.append(row_text)

    combined = "\n\n".join(full_text)

    # Use smart chunker
    smart_chunks = smart_chunk_text(combined, chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)

    chunks = []
    for i, sc in enumerate(smart_chunks):
        chunks.append(Chunk(
            text=sc.text,
            source_file=path.name,
            source_type="docx",
            chunk_index=i,
            location=f"section {i+1}",
        ))

    print(f"  [DOCX] {path.name}: {len(chunks)} chunks")

    return ExtractionResult(
        source_file=path.name,
        source_type="docx",
        chunks=chunks,
    )


# ═══════════════════════════════════════════════════════════
#  Tabular Extractor (XLSX / CSV)
# ═══════════════════════════════════════════════════════════

def _extract_tabular(path: Path, file_type: str, extracted_dir: str) -> ExtractionResult:
    rows_text = []

    if file_type == "csv":
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            reader = csv.reader(f)
            headers = None
            for i, row in enumerate(reader):
                if i == 0:
                    headers = row
                    continue
                if headers:
                    row_str = " | ".join(f"{h}: {v}" for h, v in zip(headers, row) if v.strip())
                else:
                    row_str = " | ".join(v for v in row if v.strip())
                if row_str.strip():
                    rows_text.append(row_str)
    else:
        # XLSX
        from openpyxl import load_workbook
        wb = load_workbook(str(path), read_only=True, data_only=True)
        for sheet in wb.worksheets:
            headers = None
            for i, row in enumerate(sheet.iter_rows(values_only=True)):
                vals = [str(v).strip() if v is not None else "" for v in row]
                if i == 0:
                    headers = vals
                    continue
                if headers:
                    row_str = " | ".join(f"{h}: {v}" for h, v in zip(headers, vals) if v)
                else:
                    row_str = " | ".join(v for v in vals if v)
                if row_str.strip():
                    rows_text.append(row_str)
        wb.close()

    # Group rows into chunks
    combined = "\n".join(rows_text)
    smart_chunks = smart_chunk_text(combined, chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)

    chunks = []
    for i, sc in enumerate(smart_chunks):
        chunks.append(Chunk(
            text=sc.text,
            source_file=path.name,
            source_type=file_type,
            chunk_index=i,
        ))

    print(f"  [{file_type.upper()}] {path.name}: {len(chunks)} chunks from {len(rows_text)} rows")

    return ExtractionResult(
        source_file=path.name,
        source_type=file_type,
        chunks=chunks,
    )


# ═══════════════════════════════════════════════════════════
#  Image Extractor
# ═══════════════════════════════════════════════════════════

def _extract_image(path: Path, extracted_dir: str) -> ExtractionResult:
    from PIL import Image as PILImage

    img = PILImage.open(str(path))
    width, height = img.size
    img.close()

    images = [Image(
        path=str(path),
        source_file=path.name,
        source_type="image",
        width=width,
        height=height,
    )]

    print(f"  [IMG] {path.name}: {width}x{height}")

    return ExtractionResult(
        source_file=path.name,
        source_type="image",
        images=images,
    )


# ═══════════════════════════════════════════════════════════
#  Helpers
# ═══════════════════════════════════════════════════════════

def _chunk_pages(page_texts: dict, source_file: str, source_type: str) -> list:
    """
    Chunk text from a page-keyed dict, preserving page numbers.
    Groups nearby pages when text is short, splits long pages.
    """
    chunks = []
    chunk_idx = 0

    # Accumulate text across pages, flush when big enough
    buffer = ""
    buffer_start_page = 1

    for page_num in sorted(page_texts.keys()):
        text = page_texts[page_num].strip()
        if not text:
            continue

        if not buffer:
            buffer = text
            buffer_start_page = page_num
        else:
            candidate = buffer + "\n\n" + text
            if len(candidate) > CHUNK_SIZE:
                # Flush buffer as chunk(s)
                page_chunks = _split_to_chunks(
                    buffer, source_file, source_type,
                    buffer_start_page, chunk_idx
                )
                chunks.extend(page_chunks)
                chunk_idx += len(page_chunks)
                buffer = text
                buffer_start_page = page_num
            else:
                buffer = candidate

    # Flush remaining
    if buffer.strip():
        page_chunks = _split_to_chunks(
            buffer, source_file, source_type,
            buffer_start_page, chunk_idx
        )
        chunks.extend(page_chunks)

    return chunks


def _split_to_chunks(text: str, source_file: str, source_type: str,
                     page: int, start_idx: int) -> list:
    """Split a block of text into sized chunks using the smart chunker."""
    smart_chunks = smart_chunk_text(text, chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)

    chunks = []
    for i, sc in enumerate(smart_chunks):
        if not sc.text.strip():
            continue
        chunks.append(Chunk(
            text=sc.text,
            source_file=source_file,
            source_type=source_type,
            chunk_index=start_idx + i,
            page=page,
            location=f"p.{page}",
        ))

    return chunks


def _ocr_pdf_pages(doc, filename: str) -> dict:
    """
    OCR fallback for scanned PDFs.
    Renders each page to an image, then runs Tesseract OCR to extract text.
    Returns a dict of {page_num: text}.
    """
    try:
        import pytesseract
        from PIL import Image as PILImage
        import io
    except ImportError:
        print(f"  [OCR] pytesseract not installed, skipping OCR for {filename}")
        return {}

    page_texts = {}

    for page_num in range(doc.page_count):
        try:
            page = doc[page_num]
            # Render page at 200 DPI (good balance of quality vs speed)
            import fitz as _fitz
            mat = _fitz.Matrix(200 / 72, 200 / 72)
            pix = page.get_pixmap(matrix=mat)

            # Convert to PIL Image
            img_bytes = pix.tobytes("png")
            pil_img = PILImage.open(io.BytesIO(img_bytes))

            # Run OCR
            text = pytesseract.image_to_string(pil_img).strip()

            if text and len(text) > 20:  # Skip pages with trivial OCR output
                page_texts[page_num + 1] = text

        except Exception as e:
            print(f"  [OCR] Failed on page {page_num + 1} of {filename}: {e}")
            continue

    return page_texts