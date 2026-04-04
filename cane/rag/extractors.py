"""
extractors.py — File extraction pipeline for Cane.

Supported formats:
  - PDF  (via PyMuPDF/fitz)
  - DOCX (via python-docx)
  - XLSX / CSV (via openpyxl / csv)
  - Images (OCR placeholder — stores as image for CLIP)
  - Audio / Video (faster-whisper transcription + ffmpeg keyframes)

Each extractor returns an ExtractionResult with chunks and/or images.
"""
import csv
import re
from pathlib import Path
from typing import Optional

from cane.rag.extraction_models import ExtractionResult, Chunk, Image, make_slug
from cane.rag.chunker import smart_chunk_text, smart_chunk_transcript
from cane.core.config import CHUNK_SIZE, CHUNK_OVERLAP, EXT_MAP, WHISPER_MODEL, MIN_FRAME_GAP_SEC


def extract(filepath: str, extracted_dir: str) -> ExtractionResult:
    """
    Main entry point. Detect file type and delegate to the right extractor.

    Args:
        filepath: Path to the source file
        extracted_dir: Directory for extracted images/frames

    Returns:
        ExtractionResult with chunks and images
    """
    path = Path(filepath)
    ext = path.suffix.lower()
    file_type = EXT_MAP.get(ext, "")

    if not path.exists():
        return ExtractionResult(
            source_file=path.name, source_type=file_type,
            error=f"File not found: {filepath}"
        )

    try:
        if file_type == "pdf":
            return _extract_pdf(path, extracted_dir)
        elif file_type == "docx":
            return _extract_docx(path, extracted_dir)
        elif file_type in ("xlsx", "csv"):
            return _extract_tabular(path, file_type, extracted_dir)
        elif file_type == "image":
            return _extract_image(path, extracted_dir)
        elif file_type in ("audio", "video"):
            return _extract_audio_video(path, file_type, extracted_dir)
        else:
            return ExtractionResult(
                source_file=path.name, source_type=ext,
                error=f"Unsupported file type: {ext}"
            )
    except Exception as e:
        return ExtractionResult(
            source_file=path.name, source_type=file_type,
            error=f"Extraction failed: {str(e)}"
        )


# ═══════════════════════════════════════════════════════════
#  PDF Extractor (PyMuPDF)
# ═══════════════════════════════════════════════════════════

def _extract_pdf(path: Path, extracted_dir: str) -> ExtractionResult:
    import fitz  # PyMuPDF

    doc = fitz.open(str(path))
    slug = make_slug(path.name)
    img_dir = Path(extracted_dir) / slug / "images"
    img_dir.mkdir(parents=True, exist_ok=True)

    all_text = []
    images = []
    page_texts = {}

    for page_num, page in enumerate(doc, start=1):
        # Extract text
        text = page.get_text("text").strip()
        if text:
            page_texts[page_num] = text
            all_text.append(text)

        # Extract images
        for img_idx, img_info in enumerate(page.get_images(full=True)):
            try:
                xref = img_info[0]
                base_image = doc.extract_image(xref)
                if not base_image:
                    continue

                img_bytes = base_image["image"]
                img_ext = base_image.get("ext", "png")
                width = base_image.get("width", 0)
                height = base_image.get("height", 0)

                # Skip tiny images (icons, bullets, etc.)
                if width < 100 or height < 100:
                    continue

                img_filename = f"p{page_num}_{img_idx}.{img_ext}"
                img_path = img_dir / img_filename
                with open(img_path, "wb") as f:
                    f.write(img_bytes)

                images.append(Image(
                    path=str(img_path),
                    source_file=path.name,
                    source_type="pdf",
                    page=page_num,
                    width=width,
                    height=height,
                ))
            except Exception:
                continue

    # OCR fallback: if no text was extracted but PDF has pages, it's likely scanned
    if not page_texts and doc.page_count > 0:
        print(f"  [OCR] No text found in {path.name}, running OCR on {doc.page_count} pages...")
        page_texts = _ocr_pdf_pages(doc, path.name)
        if page_texts:
            print(f"  [OCR] Extracted text from {len(page_texts)} pages")
        else:
            print(f"  [OCR] No text recovered from OCR")

    doc.close()

    # Chunk the text with page awareness
    chunks = _chunk_pages(page_texts, path.name, "pdf")

    print(f"  [PDF] {path.name}: {len(chunks)} chunks, {len(images)} images from {len(page_texts)} pages")

    return ExtractionResult(
        source_file=path.name,
        source_type="pdf",
        chunks=chunks,
        images=images,
    )


# ═══════════════════════════════════════════════════════════
#  DOCX Extractor
# ═══════════════════════════════════════════════════════════

def _extract_docx(path: Path, extracted_dir: str) -> ExtractionResult:
    from docx import Document

    doc = Document(str(path))
    full_text = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            full_text.append(text)

    # Also get tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                full_text.append(row_text)

    combined = "\n\n".join(full_text)

    # Use smart chunker
    smart_chunks = smart_chunk_text(combined, chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)

    chunks = []
    for i, sc in enumerate(smart_chunks):
        chunks.append(Chunk(
            text=sc.text,
            source_file=path.name,
            source_type="docx",
            chunk_index=i,
            location=f"section {i+1}",
        ))

    print(f"  [DOCX] {path.name}: {len(chunks)} chunks")

    return ExtractionResult(
        source_file=path.name,
        source_type="docx",
        chunks=chunks,
    )


# ═══════════════════════════════════════════════════════════
#  Tabular Extractor (XLSX / CSV)
# ═══════════════════════════════════════════════════════════

def _extract_tabular(path: Path, file_type: str, extracted_dir: str) -> ExtractionResult:
    rows_text = []

    if file_type == "csv":
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            reader = csv.reader(f)
            headers = None
            for i, row in enumerate(reader):
                if i == 0:
                    headers = row
                    continue
                if headers:
                    row_str = " | ".join(f"{h}: {v}" for h, v in zip(headers, row) if v.strip())
                else:
                    row_str = " | ".join(v for v in row if v.strip())
                if row_str.strip():
                    rows_text.append(row_str)
    else:
        # XLSX
        from openpyxl import load_workbook
        wb = load_workbook(str(path), read_only=True, data_only=True)
        for sheet in wb.worksheets:
            headers = None
            for i, row in enumerate(sheet.iter_rows(values_only=True)):
                vals = [str(v).strip() if v is not None else "" for v in row]
                if i == 0:
                    headers = vals
                    continue
                if headers:
                    row_str = " | ".join(f"{h}: {v}" for h, v in zip(headers, vals) if v)
                else:
                    row_str = " | ".join(v for v in vals if v)
                if row_str.strip():
                    rows_text.append(row_str)
        wb.close()

    # Group rows into chunks
    combined = "\n".join(rows_text)
    smart_chunks = smart_chunk_text(combined, chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)

    chunks = []
    for i, sc in enumerate(smart_chunks):
        chunks.append(Chunk(
            text=sc.text,
            source_file=path.name,
            source_type=file_type,
            chunk_index=i,
        ))

    print(f"  [{file_type.upper()}] {path.name}: {len(chunks)} chunks from {len(rows_text)} rows")

    return ExtractionResult(
        source_file=path.name,
        source_type=file_type,
        chunks=chunks,
    )


# ═══════════════════════════════════════════════════════════
#  Image Extractor
# ═══════════════════════════════════════════════════════════

def _extract_image(path: Path, extracted_dir: str) -> ExtractionResult:
    from PIL import Image as PILImage

    img = PILImage.open(str(path))
    width, height = img.size
    img.close()

    images = [Image(
        path=str(path),
        source_file=path.name,
        source_type="image",
        width=width,
        height=height,
    )]

    print(f"  [IMG] {path.name}: {width}x{height}")

    return ExtractionResult(
        source_file=path.name,
        source_type="image",
        images=images,
    )


# ═══════════════════════════════════════════════════════════
#  Audio / Video Extractor (faster-whisper + ffmpeg)
# ═══════════════════════════════════════════════════════════

def _extract_audio_video(path: Path, file_type: str, extracted_dir: str) -> ExtractionResult:
    """
    Transcribe audio/video with faster-whisper, extract keyframes from video.
    """
    import subprocess
    import tempfile

    source_file = path.name
    slug = make_slug(source_file)

    # ── Step 1: Get audio path (extract from video if needed) ──
    audio_path = str(path)
    tmp_audio = None

    if file_type == "video":
        # Extract audio track to temp WAV for whisper
        tmp_audio = tempfile.NamedTemporaryFile(suffix=".wav", delete=False)
        tmp_audio.close()
        audio_path = tmp_audio.name

        print(f"  [AV] Extracting audio from {source_file}...")
        try:
            subprocess.run(
                ["ffmpeg", "-i", str(path), "-vn", "-acodec", "pcm_s16le",
                 "-ar", "16000", "-ac", "1", audio_path, "-y"],
                capture_output=True, check=True, timeout=120,
            )
        except subprocess.CalledProcessError as e:
            _cleanup_tmp(tmp_audio)
            return ExtractionResult(
                source_file=source_file, source_type=file_type,
                error=f"ffmpeg audio extraction failed: {e.stderr.decode()[:200]}",
            )
        except FileNotFoundError:
            _cleanup_tmp(tmp_audio)
            return ExtractionResult(
                source_file=source_file, source_type=file_type,
                error="ffmpeg not installed — required for video processing",
            )

    # ── Step 2: Transcribe with faster-whisper ──
    print(f"  [AV] Transcribing with whisper ({WHISPER_MODEL})...")
    try:
        from faster_whisper import WhisperModel

        model = WhisperModel(WHISPER_MODEL, device="cpu", compute_type="int8")
        segments_iter, info = model.transcribe(audio_path, beam_size=5)

        # Collect segments
        whisper_segments = []
        for seg in segments_iter:
            whisper_segments.append({
                "text": seg.text.strip(),
                "start": seg.start,
                "end": seg.end,
            })

        # Free model memory
        del model

        print(f"  [AV] Transcribed: {len(whisper_segments)} segments, "
              f"language={info.language} ({info.language_probability:.0%})")

    except ImportError:
        _cleanup_tmp(tmp_audio)
        return ExtractionResult(
            source_file=source_file, source_type=file_type,
            error="faster-whisper not installed — required for audio/video processing",
        )
    except Exception as e:
        _cleanup_tmp(tmp_audio)
        return ExtractionResult(
            source_file=source_file, source_type=file_type,
            error=f"Transcription failed: {str(e)[:200]}",
        )
    finally:
        _cleanup_tmp(tmp_audio)

    if not whisper_segments:
        return ExtractionResult(
            source_file=source_file, source_type=file_type,
            error="No speech detected in file",
        )

    # ── Step 3: Chunk the transcript ──
    smart_chunks = smart_chunk_transcript(
        whisper_segments,
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
    )

    chunks = []
    for i, sc in enumerate(smart_chunks):
        if not sc.text.strip():
            continue

        start = sc.meta.timestamp_start or 0.0
        end = sc.meta.timestamp_end or 0.0
        location = f"{_fmt_time(start)} → {_fmt_time(end)}"

        chunks.append(Chunk(
            text=sc.text,
            source_file=source_file,
            source_type=file_type,
            chunk_index=i,
            start_sec=start,
            end_sec=end,
            location=location,
        ))

    print(f"  [AV] {len(chunks)} chunks from transcript")

    # ── Step 4: Extract keyframes from video ──
    images = []
    if file_type == "video":
        images = _extract_keyframes(path, slug, extracted_dir)

    # ── Step 5: Create OCR chunks from keyframes with slide text ──
    #    These get their own embeddings with correct text (e.g. "Phenol")
    #    so text search finds them even if Whisper mangled the transcript.
    ocr_chunks = []
    if images:
        for img in images:
            if not img.caption or len(img.caption.strip()) < 30:
                continue
            ocr_chunks.append(Chunk(
                text=img.caption.strip(),
                source_file=source_file,
                source_type=file_type,
                chunk_index=len(chunks) + len(ocr_chunks),
                start_sec=img.timestamp_sec,
                end_sec=img.timestamp_sec + MIN_FRAME_GAP_SEC,
                location=f"{_fmt_time(img.timestamp_sec)} [slide]",
            ))
        if ocr_chunks:
            print(f"  [AV] {len(ocr_chunks)} additional chunks from keyframe OCR")
            chunks.extend(ocr_chunks)

    return ExtractionResult(
        source_file=source_file,
        source_type=file_type,
        chunks=chunks,
        images=images,
    )


def _extract_keyframes(path: Path, slug: str, extracted_dir: str) -> list:
    """Extract keyframes from video at regular intervals using ffmpeg, then OCR each frame."""
    import subprocess

    img_dir = Path(extracted_dir) / slug / "images"
    img_dir.mkdir(parents=True, exist_ok=True)

    # Get video duration
    try:
        probe = subprocess.run(
            ["ffprobe", "-v", "quiet", "-print_format", "json",
             "-show_format", str(path)],
            capture_output=True, check=True, timeout=30,
        )
        import json
        duration = float(json.loads(probe.stdout).get("format", {}).get("duration", 0))
    except Exception:
        duration = 0

    if duration <= 0:
        print(f"  [AV] Could not determine video duration, skipping keyframes")
        return []

    # Extract one frame every MIN_FRAME_GAP_SEC seconds
    interval = max(MIN_FRAME_GAP_SEC, 1)
    frame_pattern = str(img_dir / f"frame_%04d.jpg")

    print(f"  [AV] Extracting keyframes every {interval}s from {duration:.0f}s video...")
    try:
        subprocess.run(
            ["ffmpeg", "-i", str(path), "-vf", f"fps=1/{interval}",
             "-qscale:v", "3", frame_pattern, "-y"],
            capture_output=True, check=True, timeout=300,
        )
    except Exception as e:
        print(f"  [AV] Keyframe extraction failed: {e}")
        return []

    # Check if OCR is available
    ocr_available = False
    try:
        import pytesseract
        pytesseract.get_tesseract_version()
        ocr_available = True
    except Exception:
        print(f"  [AV] pytesseract not available, skipping keyframe OCR")

    # Collect extracted frames + OCR each one
    images = []
    ocr_count = 0
    for frame_path in sorted(img_dir.glob("frame_*.jpg")):
        try:
            frame_num = int(frame_path.stem.split("_")[1])
            timestamp = (frame_num - 1) * interval  # ffmpeg starts at 1

            from PIL import Image as PILImage
            img = PILImage.open(str(frame_path))
            w, h = img.size

            # OCR the keyframe to extract slide text
            caption = ""
            if ocr_available:
                try:
                    ocr_text = pytesseract.image_to_string(img).strip()
                    # Only keep if meaningful text found (not just noise)
                    if ocr_text and len(ocr_text) > 20:
                        caption = ocr_text
                        ocr_count += 1
                except Exception:
                    pass

            img.close()

            images.append(Image(
                path=str(frame_path),
                source_file=path.name,
                source_type="video",
                timestamp_sec=float(timestamp),
                width=w,
                height=h,
                caption=caption,
            ))
        except Exception as e:
            print(f"  [AV] Failed reading frame {frame_path}: {e}")
            continue

    print(f"  [AV] {len(images)} keyframes extracted, {ocr_count} with OCR text")
    return images


def _fmt_time(seconds: float) -> str:
    """Format seconds as MM:SS."""
    m, s = divmod(int(seconds), 60)
    return f"{m}:{s:02d}"


def _cleanup_tmp(tmp_file):
    """Safely remove a temp file."""
    if tmp_file is not None:
        try:
            import os
            os.unlink(tmp_file.name)
        except Exception:
            pass


# ═══════════════════════════════════════════════════════════
#  Helpers
# ═══════════════════════════════════════════════════════════

def _chunk_pages(page_texts: dict, source_file: str, source_type: str) -> list:
    """
    Chunk text from a page-keyed dict, preserving page numbers.
    Groups nearby pages when text is short, splits long pages.

    Key improvements:
    - List continuation detection: if a page ends mid-list, merge with next page
    - Larger grouping: always try to keep 2-3 consecutive pages together
    - Better overlap for information-dense documents
    """
    chunks = []
    chunk_idx = 0

    # Accumulate text across pages, flush when big enough
    buffer = ""
    buffer_start_page = 1
    # Use a higher flush threshold to keep more context together
    # Flush at 1.5x chunk size so we get bigger raw blocks, then let
    # the smart chunker split them with proper overlap
    flush_threshold = int(CHUNK_SIZE * 1.5)

    sorted_pages = sorted(page_texts.keys())

    for i, page_num in enumerate(sorted_pages):
        text = page_texts[page_num].strip()
        if not text:
            continue

        if not buffer:
            buffer = text
            buffer_start_page = page_num
        else:
            candidate = buffer + "\n\n" + text
            # Check if buffer ends mid-list (bullets, numbers, dashes)
            # If so, force merge regardless of size up to 2x threshold
            ends_mid_list = _ends_mid_list(buffer)
            next_continues_list = _starts_with_list_item(text)
            force_merge = ends_mid_list or next_continues_list

            if len(candidate) > flush_threshold and not force_merge:
                # Flush buffer as chunk(s)
                page_chunks = _split_to_chunks(
                    buffer, source_file, source_type,
                    buffer_start_page, chunk_idx
                )
                chunks.extend(page_chunks)
                chunk_idx += len(page_chunks)
                buffer = text
                buffer_start_page = page_num
            elif len(candidate) > flush_threshold * 2:
                # Even with force merge, don't let it grow unbounded
                page_chunks = _split_to_chunks(
                    candidate, source_file, source_type,
                    buffer_start_page, chunk_idx
                )
                chunks.extend(page_chunks)
                chunk_idx += len(page_chunks)
                buffer = ""
                buffer_start_page = page_num + 1
            else:
                buffer = candidate

    # Flush remaining
    if buffer.strip():
        page_chunks = _split_to_chunks(
            buffer, source_file, source_type,
            buffer_start_page, chunk_idx
        )
        chunks.extend(page_chunks)

    return chunks


def _ends_mid_list(text: str) -> bool:
    """Detect if text ends in the middle of a list (bullet points, numbers, dashes)."""
    lines = text.strip().split('\n')
    if not lines:
        return False
    # Check last 3 lines for list patterns
    tail = lines[-3:] if len(lines) >= 3 else lines
    list_pattern = re.compile(r'^\s*(?:[-•●▪]\s|(?:\d+[.)]\s)|(?:[a-zA-Z][.)]\s))')
    list_lines = sum(1 for line in tail if list_pattern.match(line))
    return list_lines >= 1


def _starts_with_list_item(text: str) -> bool:
    """Detect if text starts with a list item."""
    first_line = text.strip().split('\n')[0] if text.strip() else ""
    list_pattern = re.compile(r'^\s*(?:[-•●▪]\s|(?:\d+[.)]\s)|(?:[a-zA-Z][.)]\s))')
    return bool(list_pattern.match(first_line))


def _split_to_chunks(text: str, source_file: str, source_type: str,
                     page: int, start_idx: int) -> list:
    """Split a block of text into sized chunks using the smart chunker."""
    smart_chunks = smart_chunk_text(text, chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)

    chunks = []
    for i, sc in enumerate(smart_chunks):
        if not sc.text.strip():
            continue
        chunks.append(Chunk(
            text=sc.text,
            source_file=source_file,
            source_type=source_type,
            chunk_index=start_idx + i,
            page=page,
            location=f"p.{page}",
        ))

    return chunks


def _ocr_pdf_pages(doc, filename: str) -> dict:
    """
    OCR fallback for scanned PDFs.
    Renders each page to an image, then runs Tesseract OCR to extract text.
    Returns a dict of {page_num: text}.
    """
    try:
        import pytesseract
        from PIL import Image as PILImage
        import io
    except ImportError:
        print(f"  [OCR] pytesseract not installed, skipping OCR for {filename}")
        return {}

    page_texts = {}

    for page_num in range(doc.page_count):
        try:
            page = doc[page_num]
            # Render page at 200 DPI (good balance of quality vs speed)
            import fitz as _fitz
            mat = _fitz.Matrix(200 / 72, 200 / 72)
            pix = page.get_pixmap(matrix=mat)

            # Convert to PIL Image
            img_bytes = pix.tobytes("png")
            pil_img = PILImage.open(io.BytesIO(img_bytes))

            # Run OCR
            text = pytesseract.image_to_string(pil_img).strip()

            if text and len(text) > 20:  # Skip pages with trivial OCR output
                page_texts[page_num + 1] = text

        except Exception as e:
            print(f"  [OCR] Failed on page {page_num + 1} of {filename}: {e}")
            continue

    return page_texts